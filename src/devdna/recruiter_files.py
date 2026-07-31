import csv
import re
from io import BytesIO, StringIO
from pathlib import Path

from docx import Document

from devdna.schemas import USERNAME_PATTERN

USERNAME_HEADERS = {"github", "github_id", "github_username", "username"}
GITHUB_URL_PATTERN = re.compile(r"https?://(?:www\.)?github\.com/([A-Za-z0-9-]{1,39})(?:[\s/]|$)")


class RecruiterFileError(ValueError):
    pass


def header_key(value: str) -> str:
    return value.strip().lower().replace(" ", "_").replace("-", "_")


def normalize_usernames(values: list[str], maximum: int) -> list[str]:
    usernames: list[str] = []
    seen: set[str] = set()
    for raw_value in values:
        username = raw_value.strip().removeprefix("@").lower()
        if not username:
            continue
        if not USERNAME_PATTERN.fullmatch(username) or "--" in username:
            raise RecruiterFileError(f"Invalid GitHub username: {raw_value.strip()}")
        if username not in seen:
            seen.add(username)
            usernames.append(username)
        if len(usernames) > maximum:
            raise RecruiterFileError(f"A batch can contain at most {maximum} candidates")
    if not usernames:
        raise RecruiterFileError("No GitHub usernames were found")
    return usernames


def csv_usernames(content: bytes) -> list[str]:
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise RecruiterFileError("CSV files must use UTF-8 text encoding") from error
    rows = list(csv.reader(StringIO(text)))
    if not rows:
        return []
    headers = [header_key(value) for value in rows[0]]
    header_index = next(
        (index for index, value in enumerate(headers) if value in USERNAME_HEADERS), None
    )
    if header_index is not None:
        return [row[header_index] for row in rows[1:] if len(row) > header_index]
    if all(len(row) == 1 for row in rows):
        return [row[0] for row in rows]
    raise RecruiterFileError("CSV requires a github_username column or a single username column")


def docx_usernames(content: bytes) -> list[str]:
    try:
        document = Document(BytesIO(content))
    except Exception as error:
        raise RecruiterFileError("The DOCX file could not be read") from error

    values: list[str] = []
    for paragraph in document.paragraphs:
        values.extend(GITHUB_URL_PATTERN.findall(paragraph.text))
    for table in document.tables:
        if not table.rows:
            continue
        headers = [header_key(cell.text) for cell in table.rows[0].cells]
        header_index = next(
            (index for index, value in enumerate(headers) if value in USERNAME_HEADERS),
            None,
        )
        if header_index is not None:
            values.extend(
                row.cells[header_index].text
                for row in table.rows[1:]
                if len(row.cells) > header_index
            )
    return values


def parse_recruiter_file(filename: str, content: bytes, maximum: int) -> list[str]:
    extension = Path(filename).suffix.lower()
    if extension == ".csv":
        values = csv_usernames(content)
    elif extension == ".docx":
        values = docx_usernames(content)
    else:
        raise RecruiterFileError("Upload a .csv or .docx file")
    return normalize_usernames(values, maximum)
