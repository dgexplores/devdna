import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from devdna.schemas import CvAlignment, CvSkillAlignment, EvidenceSnapshot, EvidenceSource

CV_SCHEMA_VERSION = "1"
CV_ANALYZER_VERSION = "cv-github-alignment-v1"


class CvFileError(ValueError):
    pass


@dataclass(frozen=True)
class SkillRule:
    name: str
    patterns: tuple[str, ...]
    evidence_keys: tuple[str, ...]


SKILL_RULES = (
    SkillRule("Python", (r"\bpython\b",), ("python.project",)),
    SkillRule("FastAPI", (r"\bfastapi\b",), ("api.framework.fastapi",)),
    SkillRule("Django", (r"\bdjango\b",), ("api.framework.django",)),
    SkillRule("Flask", (r"\bflask\b",), ("api.framework.flask",)),
    SkillRule("Pytest", (r"\bpytest\b",), ("testing.pytest",)),
    SkillRule(
        "Database engineering",
        (r"\bpostgres(?:ql)?\b", r"\bsqlalchemy\b", r"\balembic\b"),
        ("database.tooling",),
    ),
    SkillRule("GitHub Actions", (r"\bgithub actions\b",), ("automation.github_actions",)),
    SkillRule(
        "Container delivery",
        (r"\bdocker\b", r"\bkubernetes\b"),
        ("delivery.container",),
    ),
)


def extract_cv_text(
    filename: str,
    content: bytes,
    *,
    max_pages: int,
    max_characters: int,
) -> str:
    extension = Path(filename).suffix.lower()
    if extension == ".docx":
        try:
            document = Document(BytesIO(content))
        except Exception as error:
            raise CvFileError("The DOCX CV could not be read") from error
        parts = [paragraph.text for paragraph in document.paragraphs]
        parts.extend(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
    elif extension == ".pdf":
        try:
            reader = PdfReader(BytesIO(content))
            if reader.is_encrypted or len(reader.pages) > max_pages:
                raise CvFileError(f"PDF CVs must be unencrypted and at most {max_pages} pages")
            parts = [(page.extract_text() or "") for page in reader.pages]
        except CvFileError:
            raise
        except Exception as error:
            raise CvFileError("The PDF CV could not be read") from error
    else:
        raise CvFileError("Upload a .pdf or .docx CV")

    text = "\n".join(part.strip() for part in parts if part.strip())
    if not text:
        raise CvFileError("No readable CV text was found")
    if len(text) > max_characters:
        raise CvFileError(f"CV text exceeds the {max_characters}-character limit")
    return text


def evidence_by_key(evidence: EvidenceSnapshot) -> dict[str, list[EvidenceSource]]:
    result: dict[str, list[EvidenceSource]] = {}
    for item in evidence.items:
        result.setdefault(item.key, []).extend(item.sources)
    return result


def align_cv_to_evidence(
    github_username: str,
    source_filename: str,
    cv_text: str,
    evidence: EvidenceSnapshot,
) -> CvAlignment:
    indexed_evidence = evidence_by_key(evidence)
    skills: list[CvSkillAlignment] = []
    for rule in SKILL_RULES:
        if not any(re.search(pattern, cv_text, re.IGNORECASE) for pattern in rule.patterns):
            continue
        sources = [source for key in rule.evidence_keys for source in indexed_evidence.get(key, [])]
        skills.append(
            CvSkillAlignment(
                skill=rule.name,
                status="verified" if sources else "self_reported_unverified",
                evidence_sources=sources,
            )
        )

    verified = [skill.skill for skill in skills if skill.status == "verified"]
    unverified = [skill.skill for skill in skills if skill.status == "self_reported_unverified"]
    summary = (
        "Public GitHub work verifies " + ", ".join(verified) + "."
        if verified
        else "The uploaded CV has no matched skills verified by the current GitHub analysis."
    )
    guidance = [
        "Use verified items in the profile README and link their exact repository evidence.",
        "Keep CV-only skills out of the verified section until a public project demonstrates them.",
    ]
    if unverified:
        guidance.append("Publish reviewable evidence for: " + ", ".join(unverified) + ".")
    return CvAlignment(
        schema_version=CV_SCHEMA_VERSION,
        analyzer_version=CV_ANALYZER_VERSION,
        github_username=github_username,
        source_filename=Path(source_filename).name[:255],
        skills=skills,
        suggested_summary=summary,
        guidance=guidance,
    )
