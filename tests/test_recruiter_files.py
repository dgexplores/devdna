from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from devdna.recruiter import rank_candidates
from devdna.recruiter_files import RecruiterFileError, parse_recruiter_file
from devdna.schemas import RecruiterCandidateResult


def test_parses_deduplicated_csv_username_column() -> None:
    content = b"name,github_username\nOne,Octocat\nDuplicate,octocat\nTwo,hubot\n"

    assert parse_recruiter_file("candidates.csv", content, 10) == ["octocat", "hubot"]


def test_parses_docx_urls_and_username_table() -> None:
    document = Document()
    document.add_paragraph("Portfolio: https://github.com/octocat")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "GitHub ID"
    row = table.add_row().cells
    row[0].text = "Hubot"
    row[1].text = "hubot"
    buffer = BytesIO()
    document.save(buffer)

    assert parse_recruiter_file("candidates.docx", buffer.getvalue(), 10) == [
        "octocat",
        "hubot",
    ]


def test_rejects_invalid_or_oversized_candidate_lists() -> None:
    with pytest.raises(RecruiterFileError, match="Invalid GitHub username"):
        parse_recruiter_file("candidates.csv", b"github_username\n-invalid\n", 10)

    with pytest.raises(RecruiterFileError, match="at most 1"):
        parse_recruiter_file("candidates.csv", b"octocat\nhubot\n", 1)

    with pytest.raises(RecruiterFileError, match=".csv, .docx, or .pdf"):
        parse_recruiter_file("candidates.txt", b"octocat", 10)


def test_scans_freeform_csv_for_profile_links_and_mentions() -> None:
    content = (
        b"Candidate notes\n"
        b"Jane Doe, portfolio https://github.com/janedoe and @hubot works too\n"
        b"See https://github.com/octocat for the profile\n"
    )

    assert parse_recruiter_file("candidates.csv", content, 10) == [
        "janedoe",
        "hubot",
        "octocat",
    ]


def test_pdf_route_is_wired_and_safe(tmp_path: Path) -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    pdf_path = tmp_path / "cv.pdf"
    with pdf_path.open("wb") as handle:
        writer.write(handle)

    # A blank PDF yields no usernames instead of raising.
    with pytest.raises(RecruiterFileError, match="No GitHub usernames"):
        parse_recruiter_file("cv.pdf", pdf_path.read_bytes(), 10)

    # DOCX route still extracts bare github.com mentions and @handles.
    document = Document()
    document.add_paragraph("CV of Jane — github.com/janedoe, references @hubot")
    buffer = BytesIO()
    document.save(buffer)
    docx_names = parse_recruiter_file("cv.docx", buffer.getvalue(), 10)
    assert "janedoe" in docx_names
    assert "hubot" in docx_names


def test_candidate_ranking_uses_coverage_and_keeps_pending_unranked() -> None:
    candidates = [
        RecruiterCandidateResult(
            rank=None,
            analysis_id="one",
            github_username="lower",
            status="completed",
            requirements_met=2,
            requirements_total=7,
        ),
        RecruiterCandidateResult(
            rank=None,
            analysis_id="two",
            github_username="higher",
            status="completed",
            requirements_met=5,
            requirements_total=7,
        ),
        RecruiterCandidateResult(
            rank=None,
            analysis_id="three",
            github_username="pending",
            status="queued",
        ),
    ]

    rank_candidates(candidates)

    assert [(item.github_username, item.rank) for item in candidates] == [
        ("higher", 1),
        ("lower", 2),
        ("pending", None),
    ]
